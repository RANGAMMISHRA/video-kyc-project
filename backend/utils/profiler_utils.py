# backend/utils/profiler_utils.py
from docx import Document
from pathlib import Path
from openai import OpenAI

client = OpenAI()

def generate_kyc_overview_llm(customer_info, verifier_info, extracted_answers):
    """
    Calls an LLM to create a professional, concise customer profile overview
    combining basic info, extracted Q&A, and verifier details.
    Returns the overview text.
    """
    system_msg = (
        "You are a professional KYC profiling assistant. Generate a formal, clear overview "
        "of the customer's profile based on the provided details, including occupation, income, purpose, "
        "and risk-related information. Write in complete sentences."
    )

    prompt = f"""
Customer Details:
{customer_info}

Verifier Details:
{verifier_info}

Extracted Answers:
{extracted_answers}
"""

    print("🚀 Calling OpenAI LLM to generate profile overview...")
    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[
            {"role": "system", "content": system_msg},
            {"role": "user", "content": prompt},
        ],
    )
    return response.choices[0].message.content


def generate_profiler_doc(customer_info, verifier_info, extracted_answers, overview_text, save_path: Path):
    """
    Generates a profiler .docx with overview, customer details, verifier details, and Q&A.
    """
    doc = Document()
    doc.add_heading("Customer KYC Profile Document", level=1)

    doc.add_heading("Profile Overview", level=2)
    doc.add_paragraph(overview_text)

    doc.add_heading("Customer Details", level=2)
    for k, v in customer_info.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Verifier Details", level=2)
    for k, v in verifier_info.items():
        doc.add_paragraph(f"{k}: {v}")

    doc.add_heading("Extracted Q&A", level=2)
    for q, a in extracted_answers.items():
        doc.add_paragraph(f"{q}: {a}")

    doc.save(save_path)
    print(f"✅ Profiler document generated at {save_path}")
